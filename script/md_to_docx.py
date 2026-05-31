"""
Conversor markdown -> docx usado pra gerar docs/PIPELINE.docx a partir de
docs/PIPELINE.md. Cobre o que o documento usa: headings, paragraphs com
**bold**/*italic*/`code`, tabelas pipe, listas com `-` e `1.`, code fences
(```mermaid e ```ts), blockquotes (>), regras horizontais (---), links
[texto](url).

Blocos mermaid viram um paragrafo descritivo + caixa de codigo monospace
(docx nao renderiza mermaid). A versao online em GitHub/Replit mostra os
diagramas renderizados.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


# ---------------------------------------------------------------------------
# Helpers de estilo
# ---------------------------------------------------------------------------

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"

CYAN = RGBColor(0x18, 0x9C, 0xD9)
ANIL = RGBColor(0x3D, 0x68, 0xAB)
GRAFITE = RGBColor(0x1E, 0x1E, 0x1C)
MUTED = RGBColor(0x6B, 0x70, 0x80)
CODE_BG = "F4F4F7"


def shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size="4") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def init_styles(doc: Document) -> None:
    styles = doc.styles
    # Normal
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = BODY_FONT
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = ANIL

    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = BODY_FONT
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = CYAN

    # Heading 3
    h3 = styles["Heading 3"]
    h3.font.name = BODY_FONT
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = ANIL

    # Heading 4
    h4 = styles["Heading 4"]
    h4.font.name = BODY_FONT
    h4.font.size = Pt(12)
    h4.font.bold = True


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"          # **bold**
    r"|(`[^`]+`)"               # `code`
    r"|(\*[^*\n]+\*)"           # *italic*
    r"|(\[[^\]]+\]\([^)]+\))"   # [text](url)
)


def add_inline(paragraph, text: str) -> None:
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
            # leve fundo cinza claro
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), CODE_BG)
            rPr.append(shd)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("[") and "](" in token:
            inner = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if inner:
                label, url = inner.group(1), inner.group(2)
                run = paragraph.add_run(label)
                run.font.color.rgb = ANIL
                run.underline = True
                # nao adicionamos hyperlink real (overkill); fica visual.
                _ = url  # link textual mantido no comentario; docx nativo
                # ficou simples — manter o texto azul/underline ja sinaliza.
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# Parser principal
# ---------------------------------------------------------------------------


def parse_markdown(md_text: str, doc: Document) -> None:
    lines = md_text.split("\n")
    i = 0
    n = len(lines)
    in_code = False
    code_fence_lang = None
    code_buffer: list = []
    in_table = False
    table_rows: list = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        # primeira linha = header, segunda = separadora "---|---", resto = body
        header = table_rows[0]
        body = table_rows[2:] if len(table_rows) > 2 else []
        cols = len(header)
        table = doc.add_table(rows=1 + len(body), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        # header
        for ci, cell_text in enumerate(header):
            cell = table.rows[0].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade_cell(cell, "3D68AB")
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # body
        for ri, row in enumerate(body):
            for ci in range(cols):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                cell_md = row[ci] if ci < len(row) else ""
                add_inline(p, cell_md.strip())
                set_cell_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                if ri % 2 == 1:
                    shade_cell(cell, "F7F9FC")
        doc.add_paragraph()  # espaco depois da tabela
        table_rows = []

    def flush_code():
        nonlocal code_buffer, code_fence_lang
        if not code_buffer:
            return
        lang = (code_fence_lang or "").strip().lower()
        if lang == "mermaid":
            tag = doc.add_paragraph()
            tag_run = tag.add_run("[Diagrama Mermaid — renderizado no GitHub/Replit ao visualizar PIPELINE.md]")
            tag_run.italic = True
            tag_run.font.size = Pt(9)
            tag_run.font.color.rgb = MUTED
        # caixa monospace
        block_p = doc.add_paragraph()
        block_p.paragraph_format.left_indent = Cm(0.4)
        block_p.paragraph_format.space_before = Pt(4)
        block_p.paragraph_format.space_after = Pt(8)
        # borda esquerda fina
        pPr = block_p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side, sz in (("left", "12"),):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "8")
            b.set(qn("w:color"), "189CD9" if lang == "mermaid" else "BBBBBB")
            pBdr.append(b)
        pPr.append(pBdr)
        # shade
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_BG)
        pPr.append(shd)

        joined = "\n".join(code_buffer)
        run = block_p.add_run(joined)
        run.font.name = CODE_FONT
        run.font.size = Pt(9)
        code_buffer = []
        code_fence_lang = None

    while i < n:
        line = lines[i]

        # Code fence
        m_fence = re.match(r"^```(\w*)\s*$", line)
        if m_fence:
            if in_code:
                flush_code()
                in_code = False
            else:
                if in_table:
                    flush_table()
                    in_table = False
                in_code = True
                code_fence_lang = m_fence.group(1)
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Table detection: linha com pipes E proxima e separador
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?\s*:?-{2,}", lines[i + 1]):
            if not in_table:
                in_table = True
                table_rows = []
            # essa linha e header; proxima e separador
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            sep = lines[i + 1]
            sep_cells = [c.strip() for c in sep.strip().strip("|").split("|")]
            table_rows.append(header)
            table_rows.append(sep_cells)
            i += 2
            # consume body rows
            while i < n and "|" in lines[i] and not re.match(r"^```", lines[i]):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(row)
                i += 1
            flush_table()
            in_table = False
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        m_h = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m_h:
            level = len(m_h.group(1))
            title = m_h.group(2).strip()
            p = doc.add_heading(level=level)
            add_inline(p, title)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            b = OxmlElement("w:left")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "18")
            b.set(qn("w:space"), "8")
            b.set(qn("w:color"), "3D68AB")
            pBdr.append(b)
            pPr.append(pBdr)
            p.paragraph_format.left_indent = Cm(0.4)
            run_text = line[2:].rstrip()
            r = p.add_run()
            r.italic = True
            r.font.color.rgb = MUTED
            add_inline(p, run_text)
            i += 1
            continue

        # Lista bullet
        m_bullet = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m_bullet:
            indent = len(m_bullet.group(1)) // 2
            text = m_bullet.group(2)
            p = doc.add_paragraph(style="List Bullet")
            if indent > 0:
                p.paragraph_format.left_indent = Cm(0.6 * (indent + 1))
            add_inline(p, text)
            i += 1
            continue

        # Lista numerada
        m_num = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m_num:
            indent = len(m_num.group(1)) // 2
            text = m_num.group(2)
            p = doc.add_paragraph(style="List Number")
            if indent > 0:
                p.paragraph_format.left_indent = Cm(0.6 * (indent + 1))
            add_inline(p, text)
            i += 1
            continue

        # Linha em branco -> espaco
        if not line.strip():
            i += 1
            continue

        # Paragrafo regular (pode acumular linhas seguidas em um paragrafo so)
        buf = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(r"^(#{1,4}\s|>|[-*]\s|\d+\.\s|\|.*\||```|---+\s*$)", lines[i]):
            buf.append(lines[i])
            i += 1
        para_text = " ".join(s.strip() for s in buf)
        p = doc.add_paragraph()
        add_inline(p, para_text)


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------


def main() -> int:
    repo_root = Path(__file__).resolve().parents[1]
    # Args: opcionalmente recebe src e dst.
    if len(sys.argv) >= 3:
        src = Path(sys.argv[1]).resolve()
        dst = Path(sys.argv[2]).resolve()
    else:
        src = repo_root / "docs" / "PIPELINE.md"
        dst = repo_root / "docs" / "PIPELINE.docx"
    if not src.exists():
        print(f"NAO ENCONTRADO: {src}", file=sys.stderr)
        return 1

    md_text = src.read_text(encoding="utf-8")
    doc = Document()
    init_styles(doc)

    # Capa simples
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Lightwall Orçamento")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = ANIL

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Pipeline de Análise — Documentação Técnica")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = CYAN

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run("Gerado automaticamente a partir de docs/PIPELINE.md")
    info_run.italic = True
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = MUTED

    doc.add_page_break()

    parse_markdown(md_text, doc)

    doc.save(dst)
    print(f"OK: {dst}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
